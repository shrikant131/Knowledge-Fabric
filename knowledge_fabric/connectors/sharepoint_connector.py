"""SharePoint connector.

Uses Microsoft Graph's delta query API, which is the closest thing
SharePoint has to change tracking -- it's still poll-based (Scheduler
Adapter pattern), but the delta token means each poll only returns items
that changed since the last one, rather than a full re-scan.

Requires network access to Microsoft Graph and an OAuth app registration;
not runnable in this sandbox. Implements the same four-method interface as
every other connector so it plugs into the same registry/scheduler/pipeline
unchanged.

Auth: set the token via the environment variable named in
connector_options["auth_env_var"] (default SHAREPOINT_ACCESS_TOKEN). In
production this should come from an OAuth client-credentials flow with
automatic refresh, not a long-lived static token.
"""
from __future__ import annotations

import hashlib
import os
from typing import Iterable

from knowledge_fabric.chunking.doc_chunker import split_doc_sections
from knowledge_fabric.connectors.base import SourceConnector
from knowledge_fabric.types import Chunk, ParsedDocument, RawItem

GRAPH_BASE = "https://graph.microsoft.com/v1.0"


class SharePointConnector(SourceConnector):
    def __init__(self, source_id: str, site_id: str, drive_id: str,
                 auth_env_var: str = "SHAREPOINT_ACCESS_TOKEN",
                 delta_token_path: str = "./data/sharepoint_delta_token.txt"):
        self.source_id = source_id
        self.site_id = site_id
        self.drive_id = drive_id
        self.auth_env_var = auth_env_var
        self.delta_token_path = delta_token_path

    def _session(self):
        import requests
        token = os.environ.get(self.auth_env_var)
        if not token:
            raise RuntimeError(f"Missing credentials: set {self.auth_env_var} (Graph API access token).")
        session = requests.Session()
        session.headers.update({"Authorization": f"Bearer {token}", "Accept": "application/json"})
        return session

    def _read_delta_link(self) -> str | None:
        try:
            with open(self.delta_token_path) as f:
                return f.read().strip() or None
        except FileNotFoundError:
            return None

    def _write_delta_link(self, link: str) -> None:
        import pathlib
        pathlib.Path(self.delta_token_path).parent.mkdir(parents=True, exist_ok=True)
        with open(self.delta_token_path, "w") as f:
            f.write(link)

    # ---- fetch -----------------------------------------------------
    def fetch(self) -> Iterable[RawItem]:
        session = self._session()
        url = self._read_delta_link() or (
            f"{GRAPH_BASE}/sites/{self.site_id}/drives/{self.drive_id}/root/delta"
        )

        next_link = None
        while url:
            resp = session.get(url, timeout=30)
            resp.raise_for_status()
            payload = resp.json()
            for entry in payload.get("value", []):
                if entry.get("folder") is not None or "deleted" in entry:
                    continue  # skip folders and deletions for the pilot
                name = entry.get("name", "")
                if not _is_supported(name):
                    continue
                download_url = entry.get("@microsoft.graph.downloadUrl")
                if not download_url:
                    continue
                content_resp = session.get(download_url, timeout=60)
                content_resp.raise_for_status()
                yield RawItem(
                    source_id=self.source_id,
                    item_id=entry.get("parentReference", {}).get("path", "") + "/" + name,
                    content=_extract_text(name, content_resp.content),
                    content_type="doc",
                    language=_language_for(name),
                    extra={"web_url": entry.get("webUrl"), "modified": entry.get("lastModifiedDateTime")},
                )
            next_link = payload.get("@odata.nextLink")
            delta_link = payload.get("@odata.deltaLink")
            url = next_link
            if delta_link:
                self._write_delta_link(delta_link)

    # ---- detect_delta --------------------------------------------------
    def detect_delta(self, items: Iterable[RawItem], seen_hashes: dict[str, str]) -> list[RawItem]:
        changed = []
        for item in items:
            if seen_hashes.get(item.item_id) != item.content_hash:
                changed.append(item)
        return changed

    # ---- parse ---------------------------------------------------------
    def parse(self, item: RawItem) -> ParsedDocument:
        sections = split_doc_sections(item.content)
        return ParsedDocument(
            source_id=item.source_id, item_id=item.item_id,
            content_type="doc", language=item.language,
            sections=sections,
            extra={"content_hash": item.content_hash, "web_url": item.extra.get("web_url")},
        )

    # ---- chunk -----------------------------------------------------
    def chunk(self, doc: ParsedDocument) -> list[Chunk]:
        chunks = []
        for i, (heading, text) in enumerate(doc.sections):
            if not text.strip():
                continue
            chunk_id = hashlib.sha256(f"{doc.source_id}:{doc.item_id}:{i}".encode()).hexdigest()[:16]
            chunks.append(Chunk(
                chunk_id=chunk_id, source_id=doc.source_id, item_id=doc.item_id,
                text=text, symbol=heading, language=doc.language,
                content_hash=hashlib.sha256(text.encode()).hexdigest(),
                extra={"parent_content_hash": doc.extra.get("content_hash"), "web_url": doc.extra.get("web_url")},
            ))
        return chunks


_SUPPORTED_EXT = {".docx", ".pdf", ".txt", ".md"}


def _is_supported(name: str) -> bool:
    return any(name.lower().endswith(ext) for ext in _SUPPORTED_EXT)


def _language_for(name: str) -> str:
    return name.lower().rsplit(".", 1)[-1]


def _extract_text(name: str, content_bytes: bytes) -> str:
    lower = name.lower()
    if lower.endswith(".docx"):
        return _extract_docx_text(content_bytes)
    if lower.endswith(".pdf"):
        return _extract_pdf_text(content_bytes)
    return content_bytes.decode("utf-8", errors="ignore")


def _extract_docx_text(content_bytes: bytes) -> str:
    import io
    try:
        import docx  # python-docx
        doc = docx.Document(io.BytesIO(content_bytes))
        return "\n\n".join(p.text for p in doc.paragraphs)
    except Exception:
        return ""


def _extract_pdf_text(content_bytes: bytes) -> str:
    import io
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(content_bytes))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception:
        return ""
